from PySide6.QtWidgets import QTextEdit, QToolBar, QFileDialog
from PySide6.QtGui import QAction, QTextCursor, QTextCharFormat, QFont
from PySide6.QtPrintSupport import QPrinter, QPrintDialog
from PySide6.QtPrintSupport import QPrinter, QPrintDialog
from PySide6.QtWidgets import (
    QDialog,
    QHBoxLayout,
    QMessageBox,
    QPushButton,
    QVBoxLayout,
    QWidget,
)

import re
from html.parser import HTMLParser
from docx import Document

class RichTextEditorDialog(QDialog):
    """
    Rich text editor for last-mile formatting before printing/export.
    Supports basic formatting and inserting page breaks.
    """
    def __init__(self, parent: QWidget, *, title: str, initial_html: str):
        super().__init__(parent)
        self.setWindowTitle(title)
        self.resize(1100, 800)

        outer = QVBoxLayout(self)

        # Toolbar
        self.toolbar = QToolBar()
        outer.addWidget(self.toolbar)

        # Editor
        self.editor = QTextEdit()
        self.editor.setAcceptRichText(True)
        self.editor.setHtml(initial_html or "")
        outer.addWidget(self.editor, 1)

        # Buttons
        btns = QHBoxLayout()
        btns.addStretch(1)
        self.btn_export_pdf = QPushButton("Exportera PDF…")
        self.btn_export_docx = QPushButton("Exportera DOCX…")
        self.btn_print = QPushButton("Skriv ut…")
        self.btn_close = QPushButton("Stäng")
        btns.addWidget(self.btn_export_pdf)
        btns.addWidget(self.btn_export_docx)
        btns.addWidget(self.btn_print)
        btns.addWidget(self.btn_close)
        outer.addLayout(btns)

        self.btn_close.clicked.connect(self.reject)
        self.btn_print.clicked.connect(self._print)
        self.btn_export_pdf.clicked.connect(self._export_pdf)
        self.btn_export_docx.clicked.connect(self._export_docx)

        self._build_toolbar()

    def _build_toolbar(self) -> None:
        # Bold/Italic/Underline
        act_bold = QAction("B", self)
        act_bold.setCheckable(True)
        act_bold.triggered.connect(self._toggle_bold)
        self.toolbar.addAction(act_bold)

        act_italic = QAction("I", self)
        act_italic.setCheckable(True)
        act_italic.triggered.connect(self._toggle_italic)
        self.toolbar.addAction(act_italic)

        act_underline = QAction("U", self)
        act_underline.setCheckable(True)
        act_underline.triggered.connect(self._toggle_underline)
        self.toolbar.addAction(act_underline)

        self.toolbar.addSeparator()

        # Headings
        act_h1 = QAction("H1", self)
        act_h1.triggered.connect(lambda: self._apply_heading(1))
        self.toolbar.addAction(act_h1)

        act_h2 = QAction("H2", self)
        act_h2.triggered.connect(lambda: self._apply_heading(2))
        self.toolbar.addAction(act_h2)

        act_h3 = QAction("H3", self)
        act_h3.triggered.connect(lambda: self._apply_heading(3))
        self.toolbar.addAction(act_h3)

        act_p = QAction("P", self)
        act_p.triggered.connect(self._apply_paragraph)
        self.toolbar.addAction(act_p)

        self.toolbar.addSeparator()

        # Page break
        act_break = QAction("Sidbrytning", self)
        act_break.triggered.connect(self._insert_page_break)
        self.toolbar.addAction(act_break)

    def _toggle_bold(self, checked: bool) -> None:
        fmt = QTextCharFormat()
        fmt.setFontWeight(QFont.Bold if checked else QFont.Normal)
        self._merge_format(fmt)

    def _toggle_italic(self, checked: bool) -> None:
        fmt = QTextCharFormat()
        fmt.setFontItalic(bool(checked))
        self._merge_format(fmt)

    def _toggle_underline(self, checked: bool) -> None:
        fmt = QTextCharFormat()
        fmt.setFontUnderline(bool(checked))
        self._merge_format(fmt)

    def _merge_format(self, fmt: QTextCharFormat) -> None:
        cursor = self.editor.textCursor()
        if not cursor.hasSelection():
            cursor.select(QTextCursor.WordUnderCursor)
        cursor.mergeCharFormat(fmt)
        self.editor.mergeCurrentCharFormat(fmt)

    def _apply_heading(self, level: int) -> None:
        cursor = self.editor.textCursor()
        cursor.beginEditBlock()

        char_fmt = QTextCharFormat()
        char_fmt.setFontWeight(QFont.Bold)
        size = {1: 20, 2: 16, 3: 14}.get(level, 14)
        char_fmt.setFontPointSize(size)

        cursor.select(QTextCursor.BlockUnderCursor)
        cursor.mergeCharFormat(char_fmt)

        cursor.endEditBlock()

    def _apply_paragraph(self) -> None:
        cursor = self.editor.textCursor()
        cursor.beginEditBlock()
        char_fmt = QTextCharFormat()
        char_fmt.setFontWeight(QFont.Normal)
        char_fmt.setFontPointSize(11)
        cursor.select(QTextCursor.BlockUnderCursor)
        cursor.mergeCharFormat(char_fmt)
        cursor.endEditBlock()

    def _insert_page_break(self) -> None:
        # This works well for PDF printing from Qt and is detectable for DOCX conversion below
        cursor = self.editor.textCursor()
        cursor.insertHtml("<div style='page-break-after: always;'></div><p></p>")

    # -------- Print / PDF --------
    def _print(self) -> None:
        printer = QPrinter(QPrinter.HighResolution)
        dlg = QPrintDialog(printer, self)
        dlg.setWindowTitle("Skriv ut")
        if dlg.exec() != QDialog.Accepted:
            return
        self.editor.document().print_(printer)

    def _export_pdf(self) -> None:
        path, _ = QFileDialog.getSaveFileName(self, "Exportera PDF", "document.pdf", "PDF (*.pdf)")
        if not path:
            return
        if not path.lower().endswith(".pdf"):
            path += ".pdf"

        printer = QPrinter(QPrinter.HighResolution)
        printer.setOutputFormat(QPrinter.PdfFormat)
        printer.setOutputFileName(path)
        self.editor.document().print_(printer)
        QMessageBox.information(self, "Klart", f"PDF exporterad:\n{path}")

    # -------- DOCX --------
    def _export_docx(self) -> None:
        path, _ = QFileDialog.getSaveFileName(self, "Exportera DOCX", "document.docx", "Word (*.docx)")
        if not path:
            return
        if not path.lower().endswith(".docx"):
            path += ".docx"

        html = self.editor.document().toHtml()
        try:
            self._html_to_docx(html, path)
        except Exception as e:
            QMessageBox.critical(self, "DOCX-export misslyckades", str(e))
            return

        QMessageBox.information(self, "Klart", f"DOCX exporterad:\n{path}")

    def _html_to_docx(self, html: str, out_path: str) -> None:
        """
        Minimal, robust HTML->DOCX conversion for our editor output:
        - headings (h1/h2/h3) -> docx headings
        - paragraphs -> docx paragraphs
        - b/strong, i/em, u -> run formatting
        - br -> line break
        - page breaks inserted by our editor (<div style='page-break-after: always;'>)
        """
        doc = Document()

        # Split on our page-break marker
        parts = re.split(r"page-break-after:\s*always", html, flags=re.IGNORECASE)
        for idx, part in enumerate(parts):
            self._parse_html_into_doc(doc, part)
            if idx < len(parts) - 1:
                doc.add_page_break()

        doc.save(out_path)

    def _parse_html_into_doc(self, doc: Document, html_part: str) -> None:
        parser = _DocxHTMLParser(doc)
        parser.feed(html_part)
        parser.close()


class _DocxHTMLParser(HTMLParser):
    """
    Very small HTML parser that writes into python-docx.
    Handles: p, h1/h2/h3, br, b/strong, i/em, u
    """
    def __init__(self, doc: Document):
        super().__init__()
        self.doc = doc
        self.current_paragraph = None
        self.current_run = None
        self.bold = False
        self.italic = False
        self.underline = False
        self._block_tag = None

    def handle_starttag(self, tag, attrs):
        t = tag.lower()

        if t in ("h1", "h2", "h3"):
            level = {"h1": 1, "h2": 2, "h3": 3}[t]
            self.current_paragraph = self.doc.add_heading("", level=level)
            self.current_run = None
            self._block_tag = t
            return

        if t == "p":
            self.current_paragraph = self.doc.add_paragraph("")
            self.current_run = None
            self._block_tag = "p"
            return

        if t == "br":
            if self.current_paragraph is None:
                self.current_paragraph = self.doc.add_paragraph("")
            run = self.current_paragraph.add_run()
            run.add_break()
            return

        if t in ("b", "strong"):
            self.bold = True
        elif t in ("i", "em"):
            self.italic = True
        elif t == "u":
            self.underline = True

    def handle_endtag(self, tag):
        t = tag.lower()
        if t in ("b", "strong"):
            self.bold = False
        elif t in ("i", "em"):
            self.italic = False
        elif t == "u":
            self.underline = False
        elif t in ("p", "h1", "h2", "h3"):
            self.current_paragraph = None
            self.current_run = None
            self._block_tag = None

    def handle_data(self, data):
        text = (data or "")
        if not text.strip():
            # preserve spacing lightly
            return

        if self.current_paragraph is None:
            self.current_paragraph = self.doc.add_paragraph("")
            self._block_tag = "p"

        run = self.current_paragraph.add_run(text)
        run.bold = self.bold
        run.italic = self.italic
        run.underline = self.underline